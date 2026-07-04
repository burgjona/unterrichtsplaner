"""ASUV-Export nach LASUB-Formatvorgaben (Kap. 4.2): Arial 11, Zeilenabstand 1,5,
Blocksatz, Deckblatt, Inhaltsverzeichnis, Quellen, Selbständigkeitserklärung.

build_docx / build_pdf erhalten ein normalisiertes lesson-dict, den AsuvDraft (als
dict) und den Autorennamen und liefern die Datei als bytes.
"""
import os
from datetime import date
from io import BytesIO
from typing import List

# ---------------------------------------------------------------- gemeinsame Struktur
def _phase_rows(lesson) -> List[list]:
    header = ["Phase", "Zeit", "Sozialform", "Methode / Medien", "Lehrer-/Schülertätigkeit", "Diff. (G/M/E)"]
    rows = [header]
    for p in lesson.get("phases", []):
        rows.append([
            p.get("phase_name") or "",
            (str(p.get("minutes")) + " Min.") if p.get("minutes") is not None else "",
            p.get("social_form") or "",
            " / ".join(x for x in [p.get("method"), p.get("material")] if x),
            "L: " + (p.get("teacher_activity") or "–") + "\nS: " + (p.get("student_activity") or "–"),
            p.get("gme") or "",
        ])
    return rows


def _bibox_line(lesson) -> str:
    b = lesson.get("bibox") or {}
    if b.get("werk"):
        parts = [b["werk"], b.get("seite") or "", b.get("notiz") or ""]
        return "Lehrwerk: " + " – ".join(x for x in parts if x)
    return "Keine Lehrbuch-Referenz hinterlegt."


SELB_TEXT = ("Ich versichere, dass ich den vorliegenden Unterrichtsentwurf selbstständig und nur "
             "unter Verwendung der angegebenen Quellen und Hilfsmittel angefertigt habe.")


# ---------------------------------------------------------------- Word (.docx)
def build_docx(lesson, draft, author: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    for sname, size in (("Heading 1", 14), ("Heading 2", 12), ("Title", 20)):
        try:
            st = doc.styles[sname]
            st.font.name = "Arial"
            st.font.size = Pt(size)
            st.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    def body(text):
        for para in (text or "").split("\n"):
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not text:
            doc.add_paragraph("")

    # --- Deckblatt ---
    title = doc.add_paragraph("Ausführlicher schriftlicher Unterrichtsentwurf")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    for label, val in (
        ("Thema der Stunde", lesson.get("title", "")),
        ("Fach", lesson.get("subject", "")),
        ("Klasse / Klassenstufe", str(lesson.get("grade") or "")),
        ("Stundentyp", lesson.get("lesson_type") or ""),
        ("Schule", draft.get("schule") or ""),
        ("Verfasser/in", author),
        ("Prüfer/in / Seminarleiter/in", draft.get("pruefer") or ""),
        ("Datum", draft.get("deckblatt_datum") or date.today().isoformat()),
    ):
        pp = doc.add_paragraph()
        pp.add_run(label + ": ").bold = True
        pp.add_run(val)
    doc.add_page_break()

    # --- Inhaltsverzeichnis (Word aktualisiert Feld mit F9) ---
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tp = doc.add_paragraph()
    tp.add_run("Inhaltsverzeichnis").bold = True
    p = doc.add_paragraph()
    r = p.add_run()
    for tag, attr, txt in (("w:fldChar", ("w:fldCharType", "begin"), None),
                           ("w:instrText", ("xml:space", "preserve"), 'TOC \\o "1-3" \\h \\z \\u'),
                           ("w:fldChar", ("w:fldCharType", "separate"), None),
                           ("w:t", None, "Inhaltsverzeichnis in Word über F9 aktualisieren."),
                           ("w:fldChar", ("w:fldCharType", "end"), None)):
        el = OxmlElement(tag)
        if attr:
            el.set(qn(attr[0]), attr[1])
        if txt is not None:
            el.text = txt
        r._r.append(el)
    doc.add_page_break()

    # --- Kapitel ---
    doc.add_heading("1. Bedingungsanalyse", level=1)
    doc.add_heading("1.1 Organisatorische und technische Rahmenbedingungen", level=2)
    body(draft.get("bedingung_org"))
    doc.add_heading("1.2 Lernvoraussetzungen der Schüler:innen", level=2)
    body(draft.get("bedingung_lern"))
    doc.add_heading("1.3 Einordnung in den Lernbereich", level=2)
    body(draft.get("bedingung_einordnung"))

    doc.add_heading("2. Lehr- und Lernziele", level=1)
    body(draft.get("ziele"))

    doc.add_heading("3. Sachanalyse", level=1)
    body(draft.get("sachanalyse"))
    doc.add_heading("3.1 Quellen und Fachliteratur", level=2)
    body(draft.get("quellen"))

    doc.add_heading("4. Didaktische Analyse und methodische Entscheidungen", level=1)
    doc.add_heading("4.1 Didaktische Analyse", level=2)
    body(draft.get("didaktisch"))
    doc.add_heading("4.2 Didaktische Reduktion", level=2)
    body(draft.get("reduktion"))
    doc.add_heading("4.3 Methodische Entscheidungen", level=2)
    body(draft.get("methodisch"))

    doc.add_heading("5. Verlaufsplanung", level=1)
    rows = _phase_rows(lesson)
    if len(rows) > 1:
        table = doc.add_table(rows=0, cols=len(rows[0]))
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            cells = table.add_row().cells
            for ci, val in enumerate(row):
                cells[ci].text = val
                if ri == 0:
                    cells[ci].paragraphs[0].runs and setattr(cells[ci].paragraphs[0].runs[0], "bold", True)
    else:
        body("Noch keine Phasen erfasst.")

    doc.add_heading("6. Anhang", level=1)
    body(_bibox_line(lesson))
    body(draft.get("anhang"))

    # --- Selbständigkeitserklärung ---
    doc.add_page_break()
    se = doc.add_paragraph()
    se.add_run("Selbständigkeitserklärung").bold = True
    body(SELB_TEXT)
    doc.add_paragraph("")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Ort, Datum, Unterschrift")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- PDF
_ARIAL_REG = ["/System/Library/Fonts/Supplemental/Arial.ttf", "/Library/Fonts/Arial.ttf",
              "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
              "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"]
_ARIAL_BOLD = ["/System/Library/Fonts/Supplemental/Arial Bold.ttf", "/Library/Fonts/Arial Bold.ttf",
               "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
               "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf"]


def _first(paths):
    return next((p for p in paths if os.path.exists(p)), None)


def _register_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    reg, bold = _first(_ARIAL_REG), _first(_ARIAL_BOLD)
    if reg and bold:
        try:
            pdfmetrics.registerFont(TTFont("ASUVBody", reg))
            pdfmetrics.registerFont(TTFont("ASUVBold", bold))
            return "ASUVBody", "ASUVBold"
        except Exception:
            pass
    return "Helvetica", "Helvetica-Bold"  # immer verfügbar, Arial-metrisch ähnlich


def build_pdf(lesson, draft, author: str) -> bytes:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle)
    from reportlab.platypus.tableofcontents import TableOfContents

    body_font, bold_font = _register_fonts()
    body_st = ParagraphStyle("body", fontName=body_font, fontSize=11, leading=16.5, alignment=TA_JUSTIFY)
    h1 = ParagraphStyle("h1", fontName=bold_font, fontSize=14, leading=20, spaceBefore=12, spaceAfter=6)
    h2 = ParagraphStyle("h2", fontName=bold_font, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4)
    cover_title = ParagraphStyle("ct", fontName=bold_font, fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=24)
    cover = ParagraphStyle("cv", fontName=body_font, fontSize=12, leading=20)

    class DocT(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                style = flowable.style.name
                if style == "h1":
                    self.notify("TOCEntry", (0, flowable.getPlainText(), self.page))
                elif style == "h2":
                    self.notify("TOCEntry", (1, flowable.getPlainText(), self.page))

    buf = BytesIO()
    doc = DocT(buf, pagesize=A4, topMargin=2.5 * cm, bottomMargin=2 * cm, leftMargin=2.5 * cm, rightMargin=2 * cm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame])])

    def esc(t):
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def body(text):
        out = []
        for para in (text or "–").split("\n"):
            out.append(Paragraph(esc(para) or "&nbsp;", body_st))
        return out

    story = []
    # Deckblatt
    story += [Spacer(1, 3 * cm), Paragraph("Ausführlicher schriftlicher Unterrichtsentwurf", cover_title), Spacer(1, 1 * cm)]
    for label, val in (
        ("Thema der Stunde", lesson.get("title", "")),
        ("Fach", lesson.get("subject", "")),
        ("Klasse / Klassenstufe", str(lesson.get("grade") or "")),
        ("Stundentyp", lesson.get("lesson_type") or ""),
        ("Schule", draft.get("schule") or ""),
        ("Verfasser/in", author),
        ("Prüfer/in / Seminarleiter/in", draft.get("pruefer") or ""),
        ("Datum", draft.get("deckblatt_datum") or date.today().isoformat()),
    ):
        story.append(Paragraph("<b>%s:</b> %s" % (esc(label), esc(val)), cover))
    story.append(PageBreak())

    # Inhaltsverzeichnis
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("toc1", fontName=body_font, fontSize=11, leading=18),
        ParagraphStyle("toc2", fontName=body_font, fontSize=11, leading=16, leftIndent=18),
    ]
    story += [Paragraph("Inhaltsverzeichnis", h1), toc, PageBreak()]

    def h(text, style):
        return Paragraph(esc(text), style)

    story += [h("1. Bedingungsanalyse", h1),
              h("1.1 Organisatorische und technische Rahmenbedingungen", h2)] + body(draft.get("bedingung_org")) + \
             [h("1.2 Lernvoraussetzungen der Schüler:innen", h2)] + body(draft.get("bedingung_lern")) + \
             [h("1.3 Einordnung in den Lernbereich", h2)] + body(draft.get("bedingung_einordnung"))
    story += [h("2. Lehr- und Lernziele", h1)] + body(draft.get("ziele"))
    story += [h("3. Sachanalyse", h1)] + body(draft.get("sachanalyse")) + \
             [h("3.1 Quellen und Fachliteratur", h2)] + body(draft.get("quellen"))
    story += [h("4. Didaktische Analyse und methodische Entscheidungen", h1),
              h("4.1 Didaktische Analyse", h2)] + body(draft.get("didaktisch")) + \
             [h("4.2 Didaktische Reduktion", h2)] + body(draft.get("reduktion")) + \
             [h("4.3 Methodische Entscheidungen", h2)] + body(draft.get("methodisch"))

    story.append(h("5. Verlaufsplanung", h1))
    rows = _phase_rows(lesson)
    if len(rows) > 1:
        data = [[Paragraph(esc(c).replace("\n", "<br/>"),
                           ParagraphStyle("cell", fontName=body_font, fontSize=8, leading=11)) for c in row]
                for row in rows]
        table = Table(data, colWidths=[2.2 * cm, 1.4 * cm, 2 * cm, 3.2 * cm, 4.3 * cm, 2.4 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, (0.6, 0.6, 0.6)),
            ("BACKGROUND", (0, 0), (-1, 0), (0.9, 0.95, 0.9)),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
    else:
        story += body("Noch keine Phasen erfasst.")

    story += [h("6. Anhang", h1)] + body(_bibox_line(lesson)) + body(draft.get("anhang"))
    story += [PageBreak(), h("Selbständigkeitserklärung", h1)] + body(SELB_TEXT) + \
             [Spacer(1, 1.5 * cm), Paragraph("_______________________________", body_st),
              Paragraph("Ort, Datum, Unterschrift", body_st)]

    doc.multiBuild(story)  # zweifacher Durchlauf → korrekte TOC-Seitenzahlen
    return buf.getvalue()
