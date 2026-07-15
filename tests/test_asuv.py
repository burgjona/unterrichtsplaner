"""M6: ASUV-Entwurf – Vorbefüllung, Speichern, Export (docx/pdf)."""
from io import BytesIO

from docx import Document


def _lesson_with_phase(client):
    return client.post("/api/lessons", json={
        "title": "Balladen szenisch erschließen", "subject": "Deutsch", "grade": 8,
        "lessonType": "Einführung",
        "klafki": {"gegenwart": "Alltagsbezug", "exemplarisch": "episch-lyrische Mischform"},
        "phases": [{"phaseName": "Einstieg", "minutes": 10, "socialForm": "Plenum", "method": "Hörimpuls",
                    "teacherActivity": "spielt vor", "studentActivity": "hören"}],
    }).json()


def test_prefill_and_bibox_reminder(client, auth):
    lesson = _lesson_with_phase(client)
    a = client.get(f"/api/lessons/{lesson['id']}/asuv").json()
    assert a["saved"] is False
    assert a["biboxEmpty"] is True                     # keine BiBox → Erinnerung
    assert "Klafki" in a["ziele"]                       # Vorbefüllung aus Klafki
    assert "episch-lyrische" in a["didaktisch"]
    assert "Deutsch" in a["bedingungEinordnung"]


def test_prefill_ziele_from_lernziele(client, auth):
    """M12/U7: Kap. 2 aus Lernzielen vorbefüllt, je Feinziel Phasennachweis."""
    lesson = client.post("/api/lessons", json={
        "title": "Balladen", "subject": "Deutsch", "grade": 8,
        "phases": [
            {"phaseName": "Einstieg", "minutes": 10, "socialForm": "Plenum", "method": "Hörimpuls"},
            {"phaseName": "Erarbeitung", "minutes": 20, "socialForm": "GA", "method": "Standbild"},
        ],
        "lernziele": [
            {"kind": "grob", "text": "Balladen verstehen", "bloomStufe": "Verstehen"},
            {"kind": "fein", "text": "Merkmale nennen", "bloomStufe": "Erinnern", "phaseSortOrder": 1},
            {"kind": "fein", "text": "Wirkung reflektieren", "bloomStufe": "Bewerten", "phaseSortOrder": None},
        ],
    }).json()
    a = client.get(f"/api/lessons/{lesson['id']}/asuv").json()
    assert "Balladen verstehen" in a["ziele"]
    assert "Merkmale nennen" in a["ziele"]
    assert "erreicht in Phase: Erarbeitung" in a["ziele"]   # phaseSortOrder 1 -> zweite Phase
    assert "(keiner Phase zugeordnet)" in a["ziele"]         # phaseSortOrder None
    assert "Klafki" not in a["ziele"]                        # Lernziele haben Vorrang vor Klafki-Ableitung


def test_prefill_ziele_not_overwritten(client, auth):
    """Vorhandener Nutzertext im Feld ziele wird nie durch die Lernziel-Ableitung ersetzt."""
    lesson = client.post("/api/lessons", json={
        "title": "Balladen", "subject": "Deutsch", "grade": 8,
        "lernziele": [{"kind": "grob", "text": "Ziel A"}],
    }).json()
    client.put(f"/api/lessons/{lesson['id']}/asuv", json={"ziele": "Mein eigener Zieltext."})
    a = client.get(f"/api/lessons/{lesson['id']}/asuv").json()
    assert a["ziele"] == "Mein eigener Zieltext."


def test_save_and_persist(client, auth):
    lesson = _lesson_with_phase(client)
    put = client.put(f"/api/lessons/{lesson['id']}/asuv", json={
        "sachanalyse": "Fachwissenschaftliche Darstellung der Ballade.",
        "schule": "Oberschule Stolpen", "pruefer": "Frau Muster",
        "checks": {"0": True, "3": True}}).json()
    assert put["saved"] is True
    a = client.get(f"/api/lessons/{lesson['id']}/asuv").json()
    assert a["sachanalyse"] == "Fachwissenschaftliche Darstellung der Ballade."
    assert a["schule"] == "Oberschule Stolpen"
    assert a["checks"] == {"0": True, "3": True}


def test_export_docx_format(client, auth):
    lesson = _lesson_with_phase(client)
    client.put(f"/api/lessons/{lesson['id']}/asuv", json={"sachanalyse": "Text."})
    r = client.get(f"/api/lessons/{lesson['id']}/asuv/export?format=docx")
    assert r.status_code == 200
    assert r.content[:2] == b"PK"                        # docx = zip
    assert ".docx" in r.headers["content-disposition"]
    doc = Document(BytesIO(r.content))
    texts = "\n".join(p.text for p in doc.paragraphs)
    for heading in ["1. Bedingungsanalyse", "2. Lehr- und Lernziele", "5. Verlaufsplanung",
                    "6. Anhang", "Selbständigkeitserklärung"]:
        assert heading in texts
    assert doc.styles["Normal"].font.name == "Arial"     # Formatvorgabe
    assert len(doc.tables) == 1                           # Verlaufsplanung als Tabelle


def test_export_pdf(client, auth):
    lesson = _lesson_with_phase(client)
    r = client.get(f"/api/lessons/{lesson['id']}/asuv/export?format=pdf")
    assert r.status_code == 200
    assert r.content[:4] == b"%PDF"
    assert len(r.content) > 1500
